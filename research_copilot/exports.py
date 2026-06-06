from __future__ import annotations

import re
import textwrap
from pathlib import Path

import fitz
from docx import Document

from research_copilot.config import Settings
from research_copilot.database import Database


def export_chat_markdown(settings: Settings, db: Database, chat_id: int) -> Path:
    path = settings.exports_dir / f"chat_{chat_id}.md"
    messages = db.get_messages(chat_id)
    lines = [f"# Chat {chat_id}", ""]
    for message in messages:
        lines.append(f"## {message['role'].title()}")
        lines.append(message["content"])
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def export_chat_docx(settings: Settings, db: Database, chat_id: int) -> Path:
    path = settings.exports_dir / f"chat_{chat_id}.docx"
    doc = Document()
    doc.add_heading(f"Chat {chat_id}", level=1)
    for message in db.get_messages(chat_id):
        doc.add_heading(message["role"].title(), level=2)
        doc.add_paragraph(_strip_markdown(message["content"]))
    doc.save(path)
    return path


def export_chat_pdf(settings: Settings, db: Database, chat_id: int) -> Path:
    path = settings.exports_dir / f"chat_{chat_id}.pdf"
    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)
    y = 50
    margin = 50
    line_height = 13

    def add_line(line: str) -> None:
        nonlocal page, y
        if y > 790:
            page = pdf.new_page(width=595, height=842)
            y = 50
        page.insert_text((margin, y), line, fontsize=10, fontname="helv")
        y += line_height

    add_line(f"Chat {chat_id}")
    add_line("")
    for message in db.get_messages(chat_id):
        add_line(message["role"].upper())
        for paragraph in _strip_markdown(message["content"]).splitlines():
            for line in textwrap.wrap(paragraph, width=92) or [""]:
                add_line(line)
        add_line("")
    pdf.save(path)
    pdf.close()
    return path


def _strip_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    return text
